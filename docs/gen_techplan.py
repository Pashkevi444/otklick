# -*- coding: utf-8 -*-
"""Технический план разработки и внедрения: AI-администратор для локального бизнеса."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT2 = RGBColor(0x2E, 0x74, 0xB5)
GREY = RGBColor(0x59, 0x59, 0x59)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xB0, 0x30, 0x20)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

mono = doc.styles.add_style('Mono', 1)
mono.font.name = 'Consolas'
mono.font.size = Pt(9)

def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10, align=None, white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    if align: p.alignment = align
    run = p.add_run(text); run.bold = bold; run.font.size = Pt(size)
    if white: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    elif color: run.font.color.rgb = color
    return p

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text); run.bold = True; run.font.size = Pt(17); run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'2E74B5')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text); run.bold = True; run.font.size = Pt(13); run.font.color.rgb = ACCENT2
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    return p

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text); run.bold = True; run.font.size = Pt(11); run.font.color.rgb = GREY
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    return p

def para(text, size=11, color=None, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p

def make_table(headers, rows, col_widths=None, header_fill='1F4E79'):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        shade_cell(hdr[i], header_fill); set_cell_text(hdr[i], htext, bold=True, size=10, white=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def callout(text, fill='E8F0FE', color=None):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    cell = t.rows[0].cells[0]; shade_cell(cell, fill)
    set_cell_text(cell, text, bold=True, size=11, color=color or ACCENT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def code_block(lines, fill='F2F4F7'):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    cell = t.rows[0].cells[0]; shade_cell(cell, fill); cell.text = ''
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.style = doc.styles['Mono']
        r = p.add_run(line); r.font.name = 'Consolas'; r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ============================================================
# ТИТУЛ
# ============================================================
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI-администратор для локального бизнеса'); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=ACCENT
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Технический план разработки и внедрения'); r.bold=True; r.font.size=Pt(17); r.font.color.rgb=ACCENT2
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Архитектура · Технологический стек · Безопасность · Высокая производительность'); r.italic=True; r.font.size=Pt(12); r.font.color.rgb=GREY
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Стек: PHP 8.3 / Laravel + Octane (RoadRunner) · PostgreSQL + pgvector · Redis\n'
              'Omnichannel · RAG-LLM (YandexGPT / GigaChat) · Multi-tenant SaaS\n'
              'Соответствие 152-ФЗ · Security-by-design · Low-latency'); r.font.size=Pt(11); r.font.color.rgb=GREY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\nИюнь 2026'); r.font.size=Pt(11); r.font.color.rgb=GREY
doc.add_page_break()

# ============================================================
# 1. ОБЗОР И ЦЕЛИ
# ============================================================
h1('1. Обзор продукта и технические цели')
para('AI-администратор — мультиканальный SaaS, который перехватывает все входящие обращения '
     'локального бизнеса (WhatsApp, Telegram, Avito, VK, веб-виджет, далее — телефония), '
     'отвечает на типовые вопросы на базе знаний конкретного бизнеса и записывает клиента в его CRM. '
     'Архитектура — мультитенантная: один инстанс обслуживает множество клиентов с жёсткой изоляцией данных.')

h2('Нефункциональные требования (NFR) — то, что определяет стек')
make_table(
    ['Параметр', 'Целевое значение', 'Почему важно'],
    [
        ['Приём вебхука (ack)', '< 100 мс', 'Мессенджеры ретраят/банят при медленном ответе'],
        ['Первый токен ответа LLM', '< 1.5 с', 'Ощущение «живого» диалога'],
        ['FAQ из кэша', '< 300 мс', 'Большая часть вопросов типовая'],
        ['Доступность (uptime)', '99.9%', 'Пропуск = потерянный клиент бизнеса'],
        ['Изоляция данных клиентов', '100%', 'Multi-tenant, 152-ФЗ, доверие'],
        ['Размещение данных', 'РФ', '152-ФЗ о персональных данных'],
        ['Стоимость на 1 диалог', '< 2–4 ₽', 'Маржа 80–90%'],
    ],
    col_widths=[4.5, 4.0, 6.0]
)
callout('Три приоритета этого документа: (1) высокая скорость — за счёт Laravel Octane, '
        'очередей и стриминга LLM; (2) безопасность стека — многослойная, включая защиту от '
        'prompt-injection и изоляцию тенантов; (3) соответствие 152-ФЗ — данные и LLM в РФ.')

# ============================================================
# 2. АРХИТЕКТУРА
# ============================================================
h1('2. Архитектура системы')
para('Событийно-ориентированная архитектура: быстрый приём сообщения → постановка в очередь → '
     'асинхронная обработка ядром → ответ через канал. Тяжёлые операции (LLM, вызовы CRM) — '
     'всегда асинхронно, чтобы не блокировать приём.')

code_block([
 '                         ┌─────────────────────────────┐',
 '   WhatsApp / Telegram   │   API Gateway / Nginx (TLS) │',
 '   Avito / VK / Виджет ──▶│   + WAF + rate limit         │',
 '                         └──────────────┬──────────────┘',
 '                                        ▼',
 '                         ┌─────────────────────────────┐',
 '                         │  Webhook Receiver (Octane)  │  ← ack < 100 мс',
 '                         │  верификация подписи, enqueue│',
 '                         └──────────────┬──────────────┘',
 '                                        ▼  Redis Queue (Horizon)',
 '            ┌───────────────────────────┼───────────────────────────┐',
 '            ▼                           ▼                           ▼',
 '   ┌────────────────┐         ┌──────────────────┐        ┌────────────────┐',
 '   │ Dialog Worker  │         │  RAG / LLM слой  │        │ CRM Connector  │',
 '   │ (state machine)│◀───────▶│ pgvector + LLM   │        │ YClients / 1С  │',
 '   └───────┬────────┘         └──────────────────┘        └────────────────┘',
 '           │                          ▲                          ▲',
 '           ▼                          │                          │',
 '   ┌────────────────────────────────────────────────────────────────┐',
 '   │   PostgreSQL (tenant-isolated)  ·  Redis (cache/queue)          │',
 '   └────────────────────────────────────────────────────────────────┘',
 '           ▲',
 '           │  Inertia/Vue SPA',
 '   ┌────────────────┐',
 '   │ Панель владельца│  диалоги, метрики, база знаний, биллинг',
 '   └────────────────┘',
])

h2('Ключевые компоненты')
bullet('принимает вебхуки каналов, проверяет подпись, мгновенно кладёт в очередь и отвечает 200 OK. Никакой бизнес-логики.', 'Webhook Receiver — ')
bullet('воркеры, ведущие диалог как конечный автомат (приветствие → уточнение → ответ/запись → эскалация).', 'Dialog Workers — ')
bullet('поиск по базе знаний тенанта (RAG) + вызов LLM со стримингом, защитой от инъекций и редактированием PII.', 'RAG/LLM слой — ')
bullet('адаптеры к CRM (YClients, 1С): проверка слотов, создание записи. Изолированы за интерфейсом.', 'CRM Connectors — ')
bullet('SPA для владельца: live-диалоги, метрики ROI, редактор базы знаний, биллинг.', 'Панель управления — ')

doc.add_page_break()

# ============================================================
# 3. ТЕХ СТЕК
# ============================================================
h1('3. Технологический стек')
para('Стек выбран под три критерия: (1) твоя экспертиза (PHP), (2) высокая скорость, '
     '(3) соответствие 152-ФЗ (всё размещаемо в РФ, нет обязательной зависимости от зарубежных SaaS).')

h2('Backend')
make_table(
    ['Слой', 'Технология', 'Обоснование'],
    [
        ['Язык / рантайм', 'PHP 8.3', 'Твой профиль; JIT, типизация, fibers'],
        ['Фреймворк', 'Laravel 11', 'Скорость разработки, экосистема, очереди из коробки'],
        ['App-сервер', 'Laravel Octane + RoadRunner', 'App в памяти → x3–10 к throughput vs PHP-FPM'],
        ['Очереди', 'Redis + Laravel Horizon', 'Асинхронная обработка, мониторинг воркеров'],
        ['Realtime', 'Laravel Reverb (WebSocket)', 'Live-диалоги в панели без зарубежных сервисов'],
        ['Планировщик', 'Laravel Scheduler', 'Напоминания, реактивация, отчёты'],
    ],
    col_widths=[3.5, 4.5, 6.5]
)

h2('Данные')
make_table(
    ['Назначение', 'Технология', 'Обоснование'],
    [
        ['Основная БД', 'PostgreSQL 16', 'Надёжность, JSONB, row-level security'],
        ['Векторный поиск (RAG)', 'pgvector', 'В той же БД — меньше инфраструктуры; Qdrant если вырастет'],
        ['Кэш / очереди / сессии', 'Redis 7', 'Кэш FAQ-ответов, rate limit, очереди'],
        ['Файлы (медиа)', 'S3-совместимое (Yandex/VK Cloud, РФ)', 'Голосовые, фото; шифрование at rest'],
        ['Поиск/логи', 'OpenSearch (по мере роста)', 'Аналитика диалогов, аудит'],
    ],
    col_widths=[4.0, 5.0, 5.5]
)

h2('AI / LLM слой')
make_table(
    ['Компонент', 'Технология', 'Обоснование'],
    [
        ['LLM', 'YandexGPT / GigaChat (Sber)', 'Данные в РФ, 152-ФЗ; OpenAI-прокси только для не-ПД задач'],
        ['Эмбеддинги', 'YandexGPT Embeddings / локальный bge-m3', 'Локальная модель = нет утечки текста вовне'],
        ['Vector store', 'pgvector', 'RAG по базе знаний тенанта'],
        ['Оркестрация', 'Свой слой на PHP', 'Контроль, без тяжёлых зависимостей; LangChain не обязателен'],
        ['Голос (фаза 2)', 'SaluteSpeech / Yandex SpeechKit', 'STT/TTS для телефонии, в РФ'],
    ],
    col_widths=[3.5, 5.0, 6.0]
)

h2('Каналы (интеграции)')
make_table(
    ['Канал', 'Способ', 'Приоритет'],
    [
        ['WhatsApp', 'Официальный WhatsApp Business API через провайдера (Wazzup / 360dialog)', 'MVP'],
        ['Telegram', 'Bot API (свой сервер)', 'MVP'],
        ['Веб-виджет', 'Свой JS-виджет + WebSocket', 'MVP'],
        ['Avito / VK', 'Официальные API', 'Фаза 2'],
        ['Телефония', 'SIP + STT/TTS', 'Фаза 2'],
    ],
    col_widths=[3.0, 8.5, 2.5]
)

h2('Frontend (панель владельца)')
bullet('Inertia.js + Vue 3 + TypeScript — SPA без отдельного API-слоя, быстрая разработка.')
bullet('Tailwind CSS — скорость вёрстки, адаптив под мобильный (владельцы сидят с телефона).')
bullet('Vite — мгновенная сборка.')

h2('Инфраструктура / DevOps')
make_table(
    ['Слой', 'Технология', 'Обоснование'],
    [
        ['Хостинг', 'Российский провайдер (Yandex Cloud / Selectel / VK Cloud)', '152-ФЗ, низкая латентность к LLM'],
        ['Контейнеризация', 'Docker + Docker Compose → K8s по росту', 'Воспроизводимость, простой старт'],
        ['Reverse proxy', 'Nginx / Traefik', 'TLS, rate limit, балансировка'],
        ['CI/CD', 'GitLab CI / GitHub Actions', 'Тесты, SAST, авто-деплой'],
        ['Секреты', 'Vault / SOPS + age', 'Ключи API, токены тенантов — не в коде'],
        ['Мониторинг', 'Prometheus + Grafana + Sentry', 'Метрики, алерты, ошибки'],
        ['Логи', 'Loki / OpenSearch', 'Централизованные, с маскировкой ПД'],
    ],
    col_widths=[3.0, 5.5, 5.5]
)
doc.add_page_break()

# ============================================================
# 4. ПРОИЗВОДИТЕЛЬНОСТЬ
# ============================================================
h1('4. Высокая производительность: как держим скорость')
para('Скорость в этом продукте — это (а) мгновенный приём вебхуков, чтобы каналы не банили, '
     '(б) быстрый ответ клиенту бизнеса. Узкое место — латентность LLM, поэтому архитектура '
     'строится вокруг её сокрытия.')

h2('4.1. Laravel Octane (RoadRunner)')
bullet('Приложение загружается в память один раз и держится резидентно — нет повторного бутстрапа фреймворка на каждый запрос (как у PHP-FPM).')
bullet('Прирост пропускной способности в 3–10 раз; p99-латентность падает.')
bullet('RoadRunner (Go) предпочтительнее Swoole: стабильнее, проще, меньше «протекающего» состояния между запросами.')
para('Важно: при резидентном рантайме следить за утечками состояния между запросами '
     '(статика, синглтоны) — это и перф, и безопасность. Правило: ничего пользовательского в статике.',
     italic=True, color=GREY)

h2('4.2. Всё тяжёлое — асинхронно')
bullet('Webhook Receiver только верифицирует и кладёт в очередь → ack за < 100 мс.')
bullet('LLM-вызовы и обращения к CRM выполняются воркерами Horizon, не в HTTP-цикле.')
bullet('Отдельные очереди по приоритетам: realtime-диалоги > напоминания > аналитика.')

h2('4.3. Сокрытие латентности LLM')
bullet('Стриминг ответа токен-за-токеном — клиент видит «печатает...» и первые слова за < 1.5 с.')
bullet('Кэш типовых ответов (FAQ): семантический кэш в Redis/pgvector — повторяющиеся вопросы отдаются за < 300 мс без вызова LLM.')
bullet('Intent-shortcuts: частые намерения («адрес», «график», «цена») распознаются дешёвым классификатором и отвечаются из шаблона, минуя LLM.')
bullet('Предзагрузка RAG-контекста: эмбеддинги базы знаний посчитаны заранее, в запросе — только быстрый поиск по pgvector.')

h2('4.4. Данные и кэш')
bullet('Индексы под все запросы диалогов; HNSW-индекс для pgvector.')
bullet('Кэширование базы знаний и настроек тенанта в Redis (инвалидация при изменении).')
bullet('Пулы соединений к БД (через RoadRunner/PgBouncer), keep-alive к API LLM и провайдеров.')
bullet('N+1 под контролем (eager loading), тяжёлая аналитика — на репликах/материализованных вью.')

h2('4.5. Целевые показатели и нагрузочное тестирование')
make_table(
    ['Операция', 'Цель p95', 'Как достигается'],
    [
        ['Webhook ack', '< 100 мс', 'Octane + enqueue, без логики'],
        ['FAQ из кэша', '< 300 мс', 'Семантический кэш Redis'],
        ['Первый токен LLM', '< 1.5 с', 'Стриминг + предзагрузка RAG'],
        ['Создание записи в CRM', '< 2 с', 'Асинхронно + ретраи'],
        ['Дашборд панели', '< 500 мс', 'Кэш + индексы + Inertia'],
    ],
    col_widths=[4.5, 3.0, 7.0]
)
bullet('Нагрузочное тестирование k6 / Locust перед каждым крупным релизом; бюджет латентности зафиксирован в CI.')
doc.add_page_break()

# ============================================================
# 5. БЕЗОПАСНОСТЬ
# ============================================================
h1('5. Безопасность стека (security-by-design)')
para('Безопасность рассматривается послойно: транспорт, приложение, данные/тенанты, '
     'LLM-специфичные угрозы, секреты, инфраструктура и соответствие 152-ФЗ. '
     'Угрозы для AI-продукта шире классических — добавляются prompt-injection и утечки через LLM.')

h2('5.1. Транспорт и периметр')
bullet('TLS 1.3 везде; HSTS; mTLS для внутренних сервисов.')
bullet('WAF (Nginx + ruleset / облачный) перед приложением.')
bullet('Rate limiting и anti-abuse на уровне Nginx и приложения (по IP, по тенанту, по каналу).')
bullet('Закрытый внутренний контур: БД, Redis, воркеры — только в приватной сети, без публичного доступа.')

h2('5.2. Верификация каналов')
bullet('Обязательная проверка подписи вебхуков (Telegram secret token, HMAC провайдера WhatsApp) — отклонять неподписанные.')
bullet('Allowlist IP провайдеров там, где возможно.')
bullet('Защита от replay (nonce/timestamp).')

h2('5.3. Приложение')
bullet('Аутентификация владельцев: пароль + 2FA (TOTP); сессии в Redis, ротация.')
bullet('Авторизация: RBAC (владелец / администратор / только просмотр), проверка прав на каждом действии.')
bullet('Защита от OWASP Top-10: параметризованные запросы (Eloquent), экранирование вывода (XSS), CSRF-токены, валидация всех входов (Form Requests).')
bullet('Octane-специфика: запрет хранения пользовательских данных в статике/синглтонах между запросами; сброс состояния.')
bullet('Загрузка медиа: проверка типа/размера, антивирус-скан, хранение вне webroot.')

h2('5.4. Изоляция тенантов (multi-tenant)')
callout('Критично: один клиент НИКОГДА не должен увидеть данные другого. '
        'Изоляция проверяется автоматическими тестами на каждом релизе.', fill='FDECEA', color=RED)
bullet('tenant_id во всех таблицах + глобальный scope в Eloquent (автоподстановка фильтра).')
bullet('PostgreSQL Row-Level Security (RLS) как второй рубеж на уровне БД.')
bullet('Отдельные ключи шифрования полей на тенанта (envelope encryption).')
bullet('Изоляция RAG: векторный поиск всегда в пределах tenant_id.')
bullet('Отдельный БД-пользователь приложения с минимальными правами (least privilege).')

h2('5.5. Безопасность данных (152-ФЗ)')
bullet('Все персональные данные (имена, телефоны, переписка) — на серверах в РФ.')
bullet('Шифрование at rest (диски + чувствительные поля), in transit (TLS).')
bullet('Минимизация: не храним лишнего; ретеншн-политика и авто-удаление старых диалогов.')
bullet('Согласие на обработку ПД и политика конфиденциальности; ДПА с клиентами-бизнесами.')
bullet('Право на удаление: API/процедура удаления данных субъекта.')
bullet('Шифрованные бэкапы, регулярное тестирование восстановления.')

h2('5.6. LLM-специфичная безопасность')
para('Это то, чего нет в классическом вебе и что чаще всего упускают.', italic=True, color=GREY)
bullet('Защита от prompt-injection: системный промпт и данные пользователя строго разделены; недоверенный текст (отзывы, письма) подаётся как данные, а не инструкции.')
bullet('PII-редактирование: маскирование персональных данных перед отправкой во внешнюю LLM, где применимо.')
bullet('Guardrails на вход и выход: фильтрация попыток заставить бота выдать чужие данные, системный промпт или выполнить недопустимое действие.')
bullet('Ограничение действий LLM: бот не выполняет «опасные» операции напрямую — только через валидируемые инструменты (function calling) с проверкой прав.')
bullet('Защита от утечки системного промпта и данных других тенантов через ответы.')
bullet('Лимиты на токены/стоимость на тенанта — защита от abuse и неконтролируемых трат.')
bullet('Логирование и ревью «странных» диалогов; человек в контуре для спорных случаев.')

h2('5.7. Секреты и доступы')
bullet('Токены тенантов (CRM, WhatsApp) и ключи API — в Vault/SOPS, зашифрованы, не в репозитории.')
bullet('Ротация ключей; разные ключи для dev/stage/prod.')
bullet('Принцип наименьших привилегий для всех сервисных аккаунтов.')
bullet('Аудит-лог всех действий в панели и доступов к данным.')

h2('5.8. Безопасность цепочки поставки и процесса')
bullet('SAST (Psalm/PHPStan + security-правила) и зависимостный скан (composer audit, Dependabot) в CI.')
bullet('Pinned-зависимости, проверка лицензий.')
bullet('Code review обязателен; секреты-сканер в pre-commit (gitleaks).')
bullet('Регулярный pentest/security-review перед выходом на масштаб.')
doc.add_page_break()

# ============================================================
# 6. МОДЕЛЬ ДАННЫХ
# ============================================================
h1('6. Модель данных (основные сущности)')
make_table(
    ['Сущность', 'Назначение', 'Ключевые поля'],
    [
        ['tenants', 'Клиент-бизнес', 'id, name, plan, settings, encryption_key_id'],
        ['users', 'Владельцы/админы', 'id, tenant_id, role, 2fa_secret'],
        ['channels', 'Подключённые каналы', 'id, tenant_id, type, credentials(enc)'],
        ['contacts', 'Конечные клиенты бизнеса', 'id, tenant_id, phone(enc), name(enc)'],
        ['conversations', 'Диалоги', 'id, tenant_id, channel_id, contact_id, state, status'],
        ['messages', 'Сообщения', 'id, conversation_id, role, body, meta'],
        ['knowledge_chunks', 'База знаний (RAG)', 'id, tenant_id, content, embedding(vector)'],
        ['bookings', 'Записи в CRM', 'id, tenant_id, contact_id, crm_ref, slot, status'],
        ['crm_integrations', 'Подключения к CRM', 'id, tenant_id, type, credentials(enc)'],
        ['audit_logs', 'Аудит', 'id, tenant_id, actor, action, target, ts'],
        ['usage_metrics', 'Метрики/биллинг', 'id, tenant_id, period, dialogs, tokens, bookings'],
    ],
    col_widths=[3.5, 4.5, 6.0]
)
para('(enc) — поле шифруется на уровне приложения ключом тенанта. Все таблицы несут tenant_id '
     'и попадают под RLS.', italic=True, color=GREY)

# ============================================================
# 7. AI/RAG слой
# ============================================================
h1('7. AI / RAG слой — как бот отвечает')
h2('Конвейер обработки сообщения')
numbered('Нормализация входа, определение тенанта и диалога.')
numbered('Классификация намерения (дешёвый классификатор): FAQ-шорткат / запись / эскалация / общий вопрос.')
numbered('Если шорткат — ответ из шаблона/кэша (быстро, без LLM).')
numbered('Иначе RAG: поиск релевантных чанков базы знаний тенанта по pgvector.')
numbered('Сборка промпта: системные правила + контекст базы знаний + история диалога (данные ≠ инструкции).')
numbered('Вызов LLM со стримингом; guardrails на выход.')
numbered('Если запись — вызов CRM-инструмента (проверка слота, создание), подтверждение клиенту.')
numbered('Если уверенность низкая/чувствительный случай — эскалация на живого администратора.')
numbered('Логирование, метрики, обновление состояния диалога.')

h2('База знаний')
bullet('Владелец заполняет через панель: услуги, цены, график, адрес, частые вопросы, правила.')
bullet('Чанкинг + эмбеддинги при сохранении; пересчёт только изменённого.')
bullet('Версионирование базы знаний — откат при ошибке.')

doc.add_page_break()

# ============================================================
# 8. ПЛАН РАЗРАБОТКИ
# ============================================================
h1('8. План разработки по фазам')
para('Режим парт-тайм (10–20 ч/нед). Цель — рабочий пилот у первого клиента за ~6 недель, '
     'затем продуктизация. Каждая фаза заканчивается работающим инкрементом.')

make_table(
    ['Фаза', 'Срок', 'Содержание', 'Результат'],
    [
        ['0. Каркас', 'Нед. 1', 'Laravel+Octane, Docker, CI, БД, multi-tenant скелет, секреты', 'Деплоится «hello», тенанты изолированы'],
        ['1. Канал + эхо', 'Нед. 1–2', 'Telegram + WhatsApp webhook, верификация подписи, очереди', 'Бот принимает и отвечает шаблоном'],
        ['2. База знаний + RAG', 'Нед. 2–3', 'Редактор БЗ, эмбеддинги, pgvector, конвейер ответа', 'Бот отвечает по знаниям бизнеса'],
        ['3. LLM + стриминг', 'Нед. 3–4', 'LLM, guardrails, FAQ-кэш, intent-шорткаты', 'Живые быстрые ответы, защита от инъекций'],
        ['4. Запись в CRM', 'Нед. 4–5', 'Коннектор YClients, проверка слотов, создание записи, эскалация', 'Бот реально записывает клиентов'],
        ['5. Панель + метрики', 'Нед. 5–6', 'Live-диалоги, метрики ROI, биллинг-учёт', 'Владелец видит ценность в цифрах'],
        ['6. Хардненинг', 'Нед. 6–7', 'Security-чеклист, нагрузочное тестирование, мониторинг', 'Готовность к пилоту'],
        ['7. Пилот', 'Нед. 7–10', 'Запуск у 1–2 знакомых, замер результата', 'Кейс в цифрах'],
        ['8. Продуктизация', 'Нед. 10+', 'Self-serve онбординг, шаблоны вертикали, ускорение подключения', 'Подключение клиента за часы'],
    ],
    col_widths=[3.0, 2.0, 6.0, 4.0]
)

# ============================================================
# 9. ВНЕДРЕНИЕ
# ============================================================
h1('9. Внедрение и эксплуатация')
h2('Онбординг нового клиента (целевой процесс)')
numbered('Регистрация тенанта, подключение канала (QR WhatsApp / токен Telegram).')
numbered('Импорт/заполнение базы знаний (шаблон под вертикаль ускоряет).')
numbered('Подключение CRM (OAuth/токен), маппинг услуг и слотов.')
numbered('Тестовый прогон диалогов, настройка эскалации на администратора.')
numbered('Включение на боевом трафике, мониторинг первых дней.')
para('Цель — со 2–3 фазы продуктизации онбординг занимает часы, а не дни.', italic=True, color=GREY)

h2('Надёжность и эксплуатация')
bullet('Health-checks, авто-рестарт воркеров, graceful shutdown Octane.')
bullet('Идемпотентность обработки вебхуков (защита от дублей при ретраях каналов).')
bullet('Ретраи с backoff для CRM/LLM; circuit breaker при недоступности внешних API.')
bullet('Бэкапы БД (PITR), регулярное тестовое восстановление.')
bullet('Алерты: рост ошибок, латентность LLM, очередь, расходы на токены.')
bullet('Резервный канал (Telegram), если WhatsApp-аккаунт недоступен.')

# ============================================================
# 10. ЧЕК-ЛИСТ
# ============================================================
h1('10. Чек-лист готовности к запуску')
h2('Производительность')
for t in ['Octane (RoadRunner) в проде','Все тяжёлые операции в очередях','Стриминг LLM включён',
          'Семантический FAQ-кэш работает','Индексы БД + HNSW для pgvector','Нагрузочный тест пройден, бюджет латентности в CI']:
    p = doc.add_paragraph(); p.add_run('☐  ').font.size=Pt(12); p.add_run(t); p.paragraph_format.space_after=Pt(2)
h2('Безопасность')
for t in ['TLS 1.3, WAF, rate limiting','Верификация подписи вебхуков','2FA + RBAC в панели',
          'Изоляция тенантов: scope + RLS + тесты','Шифрование ПД at rest, данные в РФ',
          'Guardrails против prompt-injection и утечек','PII-маскирование перед внешней LLM',
          'Секреты в Vault/SOPS, не в коде','SAST + composer audit в CI','Согласия на ПД, политика, ДПА (152-ФЗ)',
          'Шифрованные бэкапы + тест восстановления','Аудит-лог действий']:
    p = doc.add_paragraph(); p.add_run('☐  ').font.size=Pt(12); p.add_run(t); p.paragraph_format.space_after=Pt(2)
h2('Продукт')
for t in ['Пилот у 1–2 клиентов','Метрики ROI собираются','Эскалация на человека работает','Шаблон онбординга вертикали']:
    p = doc.add_paragraph(); p.add_run('☐  ').font.size=Pt(12); p.add_run(t); p.paragraph_format.space_after=Pt(2)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Технический план. Следующий шаг: выбрать вертикаль и целевую CRM, '
              'после чего детализируем коннектор и базу знаний. —')
r.italic=True; r.font.size=Pt(10); r.font.color.rgb=GREY

out = '/Users/pavelbalaganskij/Desktop/AI-администратор_тех-план_разработки.docx'
doc.save(out)
print('SAVED:', out)
